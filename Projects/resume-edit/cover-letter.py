# %% cell
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re
import requests
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

matter ="""


"""

job_description = """


"""

DEEPSEEK_API = ""
OPENROUTER_API_KEY = ""



def process_with_deepseek(job_description, matter):
    prompt = f"""
    Create the body content for a professional cover letter that will be inserted between "Dear Hiring Manager," and "Sincerely,".
    The cover letter should:
    - Analyze this job description's key requirements to ensure ATS passing.
    - Highlight relevant experiences from the applicant's background
    - Match key skills from the job description
    - Maintain a professional tone
    - Never invent new skills/experiences - only repurpose existing content
    - Flow naturally from the opening and conclude approppriately before the closing.
    - Return only the body content paragraphs without any opening/closing salutations.

    Job Description:
    {job_description}

    Applicant's Background (matter):
    {matter}

    """
    url = "https://api.deepseek.com/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_API}"}
    payload = {
        "model": "deepseek-reasoner",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000
    }
    response = requests.post(url, headers=headers, json=payload)
    response_data = response.json()
    if response.status_code == 200:
        return response_data["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error: {response_data}")

def process_with_openrouter(job_description, matter):
    prompt = f"""
    Create the body content for a professional cover letter that will be inserted between "Dear Hiring Manager," and "Sincerely,".
    The cover letter should:
    - Analyze this job description's key requirements to ensure ATS passing.
    - Highlight relevant experiences from the applicant's background
    - Match key skills from the job description
    - Maintain a professional tone
    - Never invent new skills/experiences - only repurpose existing content
    - Flow naturally from the opening and conclude approppriately before the closing.
    - Return only the body content paragraphs without any opening/closing salutations.

    Job Description:
    {job_description}

    Applicant's Background (matter):
    {matter}

    """
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}"}
    payload = {
        #"model": "qwen/qwq-32b:free",
        "model": "deepseek/deepseek-chat-v3-0324:free",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000,
        "provider": {
            "order": ["Targon", "Chutes", "Fireworks"]
        }
    }
    response = requests.post(url, headers=headers, json=payload)
    response_data = response.json()
    if response.status_code == 200:
        return response_data["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error: {response_data}")


def format_markdown(doc, content):
    # Configure code style
    code_style = doc.styles.add_style('Code', 1)
    code_style.font.name = 'Consolas'
    code_style.font.size = Pt(10)
    code_paragraph_format = code_style.paragraph_format
    code_paragraph_format.left_indent = Pt(12)
    code_paragraph_format.space_before = Pt(6)
    code_paragraph_format.space_after = Pt(6)


    in_code_block = False
    code_lines = []

    for line in content.strip().split('\n'):
        stripped = line.strip()

        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                p = doc.add_paragraph(style='Code')
                p.paragraph_format.space_after = Pt(0)
                p.add_run('\n'.join(code_lines))
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Handle headings
        if re.match(r'^\*\*.*\*\*$', stripped):
            text = re.sub(r'\*\*', '', stripped)
            doc.add_heading(text, level=1)
            p.paragraph_format.line_spacing = 0
            continue

        # Handle list items
        list_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if list_match:
            _, item_text = list_match.groups()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(0)
            _process_text(p, item_text)
            continue

        # Handle regular text
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _process_text(p, line)

def _process_text(paragraph, text):
    # Handle bold formatting
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if re.match(r'^\*\*.+?\*\*$', part):
            run = paragraph.add_run(part.replace('**', ''))
            run.bold = True
        else:
            paragraph.add_run(part)

def save_cover_letter(content, output_path="generated_cover_letter.docx"):
    doc = Document()

    # Configure default style
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Add sender information
    doc.add_paragraph("Name")
    doc.add_paragraph("Address")
    doc.add_paragraph("Place")

    # Add date and city (right-aligned)
    date_para = doc.add_paragraph()
    date_para.add_run(f"                                                      place, {datetime.now().strftime('%d-%m-%Y')}")
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add salutation
    doc.add_paragraph("Dear Hiring Manager,")

    # Add formatted content
    format_markdown(doc, content)

    # Add closing
    closing_para = doc.add_paragraph("Yours Sincerly")
    closing_para.paragraph_format.space_before = Pt(24)  # Add space before signature
    doc.add_paragraph("Name")

    # Save document
    doc.save(output_path)
    print(f"Generated cover letter saved to {output_path}")


def main():
    #cover_letter_content = process_with_deepseek(job_description, matter)
    cover_letter_content = process_with_openrouter(job_description, matter)
    save_cover_letter(cover_letter_content)



if __name__ == "__main__":
    main()
