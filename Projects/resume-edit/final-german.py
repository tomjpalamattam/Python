# %% cell
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_LINE_SPACING

# Replace with your OpenRouter API key
OPENROUTER_API_KEY = ""
DEEPSEEK_API = ""

def extract_table_content(doc_path):
    """
    Extracts all subprojects (headings, subheadings, and bullet points) from the Word document.

    :param doc_path: Path to the Word document.
    :return: Dictionary containing extracted content.
    """
    # Load the document
    doc = Document(doc_path)

    # Access the first table in the document
    table = doc.tables[0]

    # Extract content from each cell
    cell_0_0 = table.cell(0, 0)  # First column
    cell_0_1 = table.cell(0, 1)  # Second column

    # Function to extract text based on formatting
    def extract_formatted_text(cell):
        content = []  # Store content in chronological order
        current_heading = None
        current_subheading = None

        for paragraph in cell.paragraphs:
            combined_text = ""  # Combine text from consecutive runs with the same formatting
            is_subheading = False
            is_heading = False
            is_bullet = False

            for run in paragraph.runs:
                # Get font size (fallback to paragraph style if run font size is None)
                font_size = run.font.size
                if font_size is None:
                    # Try to get font size from the paragraph style
                    if paragraph.style and paragraph.style.font.size:
                        font_size = paragraph.style.font.size
                    else:
                        # Fallback to a default size (e.g., 11 points)
                        font_size = 11 * 12700  # Default to 11 points in twips

                # Check if the run is part of a subheading (size 11 and bold)
                if run.bold and font_size == 11 * 12700:
                    is_subheading = True
                    combined_text += run.text
                # Check if the run is part of a heading (size 13.5 and bold)
                elif run.bold and font_size == 13.5 * 12700:
                    is_heading = True
                    combined_text += run.text
                # Check for bullet points (size 10.5)
                elif font_size == 10.5 * 12700:
                    is_bullet = True
                    combined_text += run.text
                else:
                    # If the run doesn't match any formatting, treat it as part of the previous content
                    combined_text += run.text

            # After processing all runs in the paragraph, save the combined text
            if is_subheading:
                current_subheading = combined_text.strip()
                content.append({"type": "subheading", "text": current_subheading})
            elif is_heading:
                current_heading = combined_text.strip()
                content.append({"type": "heading", "text": current_heading})
            elif is_bullet:
                if current_subheading:
                    content.append({"type": "bullet", "text": combined_text.strip(), "subheading": current_subheading})
                else:
                    content.append({"type": "bullet", "text": combined_text.strip(), "heading": current_heading})

        return content

    # Extract content from both cells
    content_0_0 = extract_formatted_text(cell_0_0)
    content_0_1 = extract_formatted_text(cell_0_1)

    return {
        "cell_0_0": content_0_0,
        "cell_0_1": content_0_1,
    }

def add_bullet_point(paragraph):
    """
    Adds a bullet point to a paragraph.
    """
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), '1')  # Use the first numbering style (bullet points)
    numPr.append(numId)
    pPr.append(numPr)

def add_key_points(doc, subheading, key_points):
    """
    Adds key points to the specified subheading in the original document.
    Removes leading and trailing whitespace from each bullet point.
    """
    # Access the first table in the document
    table = doc.tables[0]

    # Iterate through the cells to find the subheading
    for row in table.rows:
        for cell in row.cells:
            # Find the subheading paragraph and the date paragraph (if it exists)
            subheading_paragraph = None
            date_paragraph = None
            for i, paragraph in enumerate(cell.paragraphs):
                if paragraph.text.strip() == subheading:
                    subheading_paragraph = paragraph
                    # Check if the next paragraph is a date
                    if i + 1 < len(cell.paragraphs):
                        next_paragraph = cell.paragraphs[i + 1]
                        # Check if the next paragraph is a date (e.g., "Oct 2023 — Jul 2024")
                        if "—" in next_paragraph.text.strip():
                            date_paragraph = next_paragraph
                    break

            # If the subheading is found, insert bullet points after the date (if it exists) or after the subheading
            if subheading_paragraph:
                # Determine the insertion point
                insertion_point = date_paragraph if date_paragraph else subheading_paragraph

                # Reverse the key_points list to ensure correct order when adding
                key_points_reversed = key_points[::-1]

                # Insert key points as bullet points after the insertion point
                for key_point in key_points_reversed:
                    # Remove leading and trailing whitespace from the bullet point
                    cleaned_key_point = key_point.strip()

                    # Skip empty bullet points
                    if not cleaned_key_point:
                        continue

                    new_paragraph = cell.add_paragraph()
                    new_paragraph.text = cleaned_key_point
                    add_bullet_point(new_paragraph)  # Add bullet point
                    # Set font to Arial and size to 10.5 points
                    for run in new_paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10.5)

                    # Move the new paragraph after the insertion point
                    insertion_point._element.addnext(new_paragraph._element)

    return doc


def process_with_openrouter(job_description, project_data):
    prompt = f"""
    ROLE: You are a strict German resume editor optimizing for ATS compliance.
    TASKS:
    1. Analyze this job description's key requirements
    2. Rewrite ONLY ACTUAL PROJECTS from provided data using German business language
    3. Never invent new skills/experiences - only repurpose existing content
    4. Keep bullets concise (12-15 words max) with quantifiable results
    5. Maximum 5-6 most relevant bullets

    RULES:
    - Output format: "bullet1|bullet2|bullet3" (NO MARKERS, ONLY TEXT)
    - No explanations or formatting
    - Match exact technical terms from job description
    - Prioritize measurable achievements
    - Allow mild exaggeration of existing elements but NO fabrication

    JOB DESCRIPTION:
    {job_description}

    ACTUAL PROJECT INPUT:
    {project_data}

    OUTPUT (GERMAN):"""

    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}"}
    payload = {
        #"model": "qwen/qwq-32b:free",
        "model": "deepseek/deepseek-chat-v3-0324:free",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000,
        "provider": {
            "order": ["Targon", "Chutes", "Fireworks"]
        }
    }

    response = requests.post(url, headers=headers, json=payload)
    response_data = response.json()

    if response.status_code == 200:
        return response_data["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error: {response_data}")


def process_with_deepseek(job_description, project_data):
    prompt = f"""
    ROLE: You are a strict German resume editor optimizing for ATS compliance.
    TASKS:
    1. Analyze this job description's key requirements
    2. Rewrite ONLY ACTUAL PROJECTS from provided data using German business language
    3. Never invent new skills/experiences - only repurpose existing content
    4. Keep bullets concise (12-15 words max) with quantifiable results
    5. Maximum 5-6 most relevant bullets

    RULES:
    - Output format: "bullet1|bullet2|bullet3" (NO MARKERS, ONLY TEXT)
    - No explanations or formatting
    - Match exact technical terms from job description
    - Prioritize measurable achievements
    - Allow mild exaggeration of existing elements but NO fabrication

    JOB DESCRIPTION:
    {job_description}

    ACTUAL PROJECT INPUT:
    {project_data}

    OUTPUT (GERMAN):"""

    url = "https://api.deepseek.com/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_API}"}
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000
    }

    response = requests.post(url, headers=headers, json=payload)
    response_data = response.json()

    if response.status_code == 200:
        return response_data["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error: {response_data}")


def process_with_openrouter_skills(job_description, project_data):
    prompt = f"""
    You are an AI that optimizes resumes for ATS scoring based on a given job description.
    Given the job description below, analyze the data and:

    rewrite the skills  (must be in german language) that is provided by either removing some or adding new ones to match the job description.
    generate the skills with '|' seperator and nothing else

    JOB DESCRIPTION:
    {job_description}

    SKILLS ():
    {project_data}

    """

    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {"Authorization": f"Bearer {OPENROUTER_API_KEY}"}
    payload = {
        #"model": "qwen/qwq-32b:free",
        "model": "deepseek/deepseek-chat-v3-0324:free",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000,
        "provider": {
            "order": ["Targon", "Chutes", "Fireworks"]
        }
    }

    response = requests.post(url, headers=headers, json=payload)
    response_data = response.json()

    if response.status_code == 200:
        return response_data["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error: {response_data}")

def process_with_deepseek_skills(job_description, project_data):
    prompt = f"""
    You are an AI that optimizes resumes for ATS scoring based on a given job description.
    Given the job description below, analyze the data and:

    rewrite the skills that is provided by either removing some or adding new ones to match the job description.
    generate the skills with '|' seperator and nothing else

    JOB DESCRIPTION:
    {job_description}

    SKILLS ():
    {project_data}

    """

    url = "https://api.deepseek.com/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_API}"}
    payload = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000
    }

    response = requests.post(url, headers=headers, json=payload)
    response_data = response.json()

    if response.status_code == 200:
        return response_data["choices"][0]["message"]["content"]
    else:
        raise Exception(f"Error: {response_data}")


def process_with_openrouter_delete_sessions(job_description, project_data):
    """Same as original implementation"""
    prompt = f"""
    You are an AI that optimizes resumes for ATS scoring based on a given job description.
    Given the job description below, analyze the data and:

    Choose whether a project is relevant or not. You are only supposed to give a boolean answer (0 or 1).

    JOB DESCRIPTION:
    {job_description}

    CURRENT PROJECT:
    {project_data}
    """

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers={"Authorization": f"Bearer {OPENROUTER_API_KEY}"},
        json={
            #"model": "qwen/qwq-32b:free",
            "model": "deepseek/deepseek-chat-v3-0324:free",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 1000,
            "provider": {
                "order": ["Targon", "Chutes", "Fireworks"]
            }
        }
    )
    if response.status_code == 200:
        return int(response.json()["choices"][0]["message"]["content"].strip())
    else:
        raise Exception(f"API Error: {response.text}")


def process_with_deepseek_delete_sessions(job_description, project_data):
    """Same as original implementation"""
    prompt = f"""
    You are an AI that optimizes resumes for ATS scoring based on a given job description.
    Given the job description below, analyze the data and:

    Choose whether a project is relevant or not. You are only supposed to give a boolean answer (0 or 1).

    JOB DESCRIPTION:
    {job_description}

    CURRENT PROJECT:
    {project_data}
    """

    url = "https://api.deepseek.com/chat/completions"
    headers = {"Authorization": f"Bearer {DEEPSEEK_API}"}
    payload = {
        #"model": "qwen/qwq-32b:free",
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": 1000
    }
    response = requests.post(url, headers=headers, json=payload)
    if response.status_code == 200:
        return int(response.json()["choices"][0]["message"]["content"].strip())
    else:
        raise Exception(f"API Error: {response.text}")


def extract_all_subprojects(content):
    """
    Extracts all subprojects from the content.

    :param content: Dictionary containing extracted content.
    :return: List of subprojects, each as a dictionary.
    """
    subprojects = []
    for cell, data in content.items():
        current_subproject = None
        for item in data:
            if item["type"] == "subheading":
                # Start a new subproject
                current_subproject = {
                    "subheading": item["text"],
                    "bullets": []
                }
                subprojects.append(current_subproject)
            elif item["type"] == "bullet" and current_subproject:
                # Add bullet points to the current subproject
                current_subproject["bullets"].append(item["text"])
    return subprojects

def extract_all_projects(content):
    """
    Extracts all projects from the content.

    :param content: Dictionary containing extracted content.
    :return: List of projects, each as a dictionary.
    """
    projects = []
    for cell, data in content.items():
        current_project = None
        for item in data:
            if item["type"] == "heading":
                # Start a new subproject
                current_project = {
                    "heading": item["text"],
                    "bullets": []
                }
                projects.append(current_project)
            elif item["type"] == "bullet" and current_project:
                # Add bullet points to the current project
                current_project["bullets"].append(item["text"])
    return projects

def update_skills_section_old(doc, skills):
    """Updates skills section with exact formatting requirements"""
    table = doc.tables[0]
    cell = table.cell(0, 1)  # cell_0_1

    # Configure paragraph formatting
    para_format = {
        'font_name': 'Arial',
        'font_size': Pt(10),
        'space_after': Pt(0),
        'line_spacing': 2.0,
        'line_spacing_rule': WD_LINE_SPACING.EXACTLY
    }

    # Find or create skills heading
    skills_heading = next((p for p in cell.paragraphs if "Fähigkeiten" in p.text), None)

    if not skills_heading:
        skills_heading = cell.add_paragraph("Fähigkeiten")
        skills_heading.runs[0].bold = True
        skills_heading.runs[0].font.name = 'Arial'
        skills_heading.runs[0].font.size = Pt(11)

    # Add skills with precise formatting
    for skill in skills:
        new_para = cell.add_paragraph()
        new_para.text = skill
        pf = new_para.paragraph_format
        pf.space_after = para_format['space_after']
        pf.line_spacing = para_format['line_spacing']
        pf.line_spacing_rule = para_format['line_spacing_rule']

        for run in new_para.runs:
            run.font.name = para_format['font_name']
            run.font.size = para_format['font_size']


def update_skills_section(doc, skills):
    """Add new skills at the beginning of skills section with precise formatting"""
    table = doc.tables[0]
    cell = table.cell(0, 1)  # cell_0_1

    # Find or create skills heading
    skills_heading = None
    for para in cell.paragraphs:
        if "Fähigkeiten" in para.text:
            skills_heading = para
            break

    if not skills_heading:
        skills_heading = cell.add_paragraph("Fähigkeiten")
        skills_heading.runs[0].bold = True
        skills_heading.runs[0].font.name = 'Arial'
        skills_heading.runs[0].font.size = Pt(11)

    # Create insertion point after heading
    insertion_point = skills_heading._element

    # Add new skills in reverse order to appear in correct order
    for skill in reversed(skills):
        new_para = cell.add_paragraph()
        new_para.text = skill.strip()

        # Set formatting
        for run in new_para.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)

        # Set paragraph spacing
        para_format = new_para.paragraph_format
        para_format.line_spacing = 2.0
        para_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para_format.space_after = Pt(0)

        # Insert after heading
        insertion_point.addnext(new_para._element)
        insertion_point = new_para._element


# Function to find and remove subheadings in template
def clean_template(template_doc, irrelevant_subheadings):
    template_table = template_doc.tables[0]
    cells_to_process = [template_table.cell(0, 0), template_table.cell(0, 1)]

    for cell in cells_to_process:
        paragraphs_to_remove = []

        # First pass: Identify paragraphs to remove
        for paragraph in cell.paragraphs:
            text = paragraph.text.strip()
            if text in irrelevant_subheadings and text not in EXCLUDED_SUBHEADINGS_DELETION:
                paragraphs_to_remove.append(paragraph)

        # Second pass: Remove identified paragraphs
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)


def filter_irrelevant_subprojects(subprojects, irrelevant_subheadings):
    """
    Filters out subprojects with irrelevant subheadings.

    :param subprojects: List of subprojects, each as a dictionary.
    :param irrelevant_subheadings: List of subheadings to be removed.
    :return: Filtered list of subprojects.
    """
    return [subproject for subproject in subprojects if subproject["subheading"] not in irrelevant_subheadings]




# Path to your CV file
cv_path = 'templates/CV-Deutsch.docx'

# Extract content from the table
content = extract_table_content(cv_path)

# Extract all subprojects
subprojects = extract_all_subprojects(content)

EXCLUDED_SUBHEADINGS_DELETION = [] 
EXCLUDED_SUBHEADINGS = []
# Job description
job_desc = """


"""
# Load the document
doc = Document("templates/CV-d-template.docx")

# Evaluate relevance with exclusion check
results = {}
irrelevant_subheadings = []

for subproject in subprojects:
    project_data = f"Subheading: {subproject['subheading']}\nBullets:\n"
    project_data += "\n".join([bullet for bullet in subproject["bullets"]])

    try:
        if subproject['subheading'] in EXCLUDED_SUBHEADINGS_DELETION:
            results[subproject["subheading"]] = 1
            continue

        relevance = process_with_openrouter_delete_sessions(job_desc, project_data)
        #relevance = process_with_deepseek_delete_sessions(job_desc, project_data)
        results[subproject["subheading"]] = relevance

        if not relevance:
            irrelevant_subheadings.append(subproject["subheading"])
    except Exception as e:
        print(f"Error processing {subproject['subheading']}: {str(e)}")

# Clean the template document
clean_template(doc, irrelevant_subheadings)

# Filter out irrelevant subprojects from the subprojects list
subprojects = filter_irrelevant_subprojects(subprojects, irrelevant_subheadings)


# Print results with protection notice
print("\nOptimization Results:")
for subheading, relevant in results.items():
    status = "🔒 Protected" if subheading in EXCLUDED_SUBHEADINGS_DELETION else '✅ Kept' if relevant else '❌ Removed'
    print(f"- {status} {subheading}")

# Rewrite Bullet Points
results = {}
for subproject in subprojects:
    subheading = subproject["subheading"]

    # Skip processing if the subheading is in the excluded list
    if subheading in EXCLUDED_SUBHEADINGS:
        print(f"Skipping excluded subheading: {subheading}")
        continue  # Move to the next subproject

    subproject_data = f"Subheading: {subheading}\n"
    subproject_data += "Bullets:\n" + "\n".join(subproject["bullets"])
    print(f"Processing data for subheading: {subheading}")

    try:
        # Generate optimized bullet points using OpenRouter
        key_points_string = process_with_openrouter(job_desc, subproject_data)
        #key_points_string = process_with_deepseek(job_desc, subproject_data)
        key_points = key_points_string.split("|")

        # Add the optimized bullet points to the document
        updated_doc = add_key_points(doc, subheading, key_points)
    except Exception as e:
        print(f"Error evaluating subproject '{subheading}': {e}")

# Updating the skills section

cell_0_1_content = content["cell_0_1"]
start = False
cell_0_1_data = [item['text'] for item in cell_0_1_content
          if (item['text'] == "Fähigkeiten" and (start := True)) or
             (start and item['text'] != "Sprachen" and not (item['text'] == "Sprachen" and False))]

new_skills = []
new_skills = process_with_openrouter_skills(job_desc,cell_0_1_data)
#new_skills = process_with_deepseek_skills(job_desc,cell_0_1_data)
key_skills = new_skills.split("|")

update_skills_section(updated_doc,key_skills)

# Save the updated document (overwrite the original file)
updated_doc.save('CV_MOD.docx')
